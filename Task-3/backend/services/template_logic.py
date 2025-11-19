from __future__ import annotations

import io
import json
from typing import Dict, Set

import docx

from backend.config import get_settings
from backend.services.llm_client import create_chat_completion


def _load_document(template_bytes: bytes) -> docx.Document:
    return docx.Document(io.BytesIO(template_bytes))


def extract_template_text(template_bytes: bytes) -> str:
    document = _load_document(template_bytes)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _heuristic_field_candidates(template_bytes: bytes) -> Set[str]:
    doc = _load_document(template_bytes)
    candidates: Set[str] = set()

    def add_candidate(text: str) -> None:
        cleaned = text.strip().strip(":")
        if 1 < len(cleaned) <= 80:
            candidates.add(cleaned)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "{{" in text and "}}" in text:
            for chunk in text.split("{{"):
                if "}}" in chunk:
                    add_candidate(chunk.split("}}", 1)[0])
        elif ":" in text:
            add_candidate(text.split(":", 1)[0])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                if ":" in cell_text:
                    add_candidate(cell_text.split(":", 1)[0])
    return candidates


def detect_fields_with_llm(template_bytes: bytes) -> Dict[str, str]:
    settings = get_settings()
    template_text = extract_template_text(template_bytes)
    candidates = sorted(_heuristic_field_candidates(template_bytes))

    prompt = f"""Extract field names from this insurance template. Return a JSON object where keys are field names and values are empty strings.

TEMPLATE TEXT:
{template_text[:settings.max_report_chars]}

CANDIDATE FIELDS:
{candidates}

Return ONLY JSON, no markdown or commentary."""

    response = create_chat_completion(
        messages=[
            {"role": "system", "content": "You extract field names from templates and return JSON."},
            {"role": "user", "content": prompt},
        ]
    )

    try:
        cleaned = response.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned.split("```json", 1)[1].split("```", 1)[0].strip()
        elif cleaned.startswith("```"):
            cleaned = cleaned.split("```", 1)[1].split("```", 1)[0].strip()
        return json.loads(cleaned)
    except json.JSONDecodeError as exc:
        print(f"[template_logic] Failed to parse LLM response as JSON: {response[:500]}")
        fallback = {label: "" for label in candidates}
        if fallback:
            return fallback
        raise ValueError("Failed to parse template fields JSON") from exc


def fill_template(template_bytes: bytes, data: Dict[str, str]) -> bytes:
    document = _load_document(template_bytes)

    # Fill paragraphs - look for "Field Name:" patterns and insert value after
    for paragraph in document.paragraphs:
        original_text = paragraph.text
        if not original_text.strip():
            continue
            
        for key, value in data.items():
            if not value:
                continue
                
            # Check if this paragraph contains the field name followed by a colon
            if f"{key}:" in original_text or f"{key} :" in original_text:
                # Replace the entire paragraph with "Field Name:\nValue"
                paragraph.clear()
                paragraph.add_run(f"{key}:\n{value}")
                break
            elif f"{{{{{key}}}}}" in original_text:
                # Handle {{placeholder}} style
                paragraph.text = original_text.replace(f"{{{{{key}}}}}", value)
                break

    # Fill tables
    for table in document.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                    
                for key, value in data.items():
                    if not value:
                        continue
                        
                    # Check if this cell contains the field name
                    if key in cell_text and ":" in cell_text:
                        # Field name with colon in same cell - replace after colon
                        cell.text = f"{key}:\n{value}"
                        break
                    elif key == cell_text and idx + 1 < len(row.cells):
                        # Field name alone in cell, value goes in next cell
                        row.cells[idx + 1].text = value
                        break

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
